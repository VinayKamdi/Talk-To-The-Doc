import streamlit as st
import os
import shutil
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain.docstore.document import Document as LangchainDocument
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from pdf2image import convert_from_path
from pytesseract import image_to_string
from docx import Document
from PIL import Image
import magic

load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

st.title("Talk to the Doc.")

llm = ChatGroq(groq_api_key=groq_api_key, model_name="Llama3-8b-8192") # type: ignore

prompt = ChatPromptTemplate.from_template(
    """
    Provided the document, give most accurate response based on question.
    <document>
    {context}
    <document>
    Questions: {input}
    """
)

def extract_text_from_doc(file_path):
    doc = Document(file_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return text

def extract_text_from_image(image_path):
    image = Image.open(image_path)
    text = image_to_string(image)
    return text

def extract_text_from_pdf_images(pdf_path):
    images = convert_from_path(pdf_path, poppler_path=r'./poppler-24.07.0/Library/bin')
    text = ''
    for image in images:
        text += image_to_string(image)
    return text

def vector_embedding(files):
    if "vectors" not in st.session_state:
        temp_dir = "./temp_files"
        os.makedirs(temp_dir, exist_ok=True)

        documents = []
        for file in files:
            file_type = magic.from_buffer(file.read(1024), mime=True)
            file.seek(0)

            file_path = os.path.join(temp_dir, file.name)
            with open(file_path, "wb") as f:
                f.write(file.getbuffer())

            if file_type == "application/pdf":
                loader = PyPDFLoader(file_path)
                documents.extend(loader.load())

                image_text = extract_text_from_pdf_images(file_path)
                documents.append(LangchainDocument(page_content=image_text))

            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc_text = extract_text_from_doc(file_path)
                documents.append(LangchainDocument(page_content=doc_text))

            elif "image" in file_type:
                image_text = extract_text_from_image(file_path)
                documents.append(LangchainDocument(page_content=image_text))

            else:
                st.warning(f"Unsupported file format: {file_type}")

        st.session_state.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001") # type: ignore
        st.session_state.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        st.session_state.final_documents = st.session_state.text_splitter.split_documents(documents)
        st.session_state.vectors = FAISS.from_documents(st.session_state.final_documents, st.session_state.embeddings)

        for file in files:
            os.remove(os.path.join(temp_dir, file.name))
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

uploaded_files = st.file_uploader("Upload documents (PDF, DOCX, Images)", type=["pdf", "docx", "jpg", "jpeg", "png"], accept_multiple_files=True)
if uploaded_files:
    vector_embedding(uploaded_files)
    prompt1 = st.text_input("Ask questions.")
    
    if prompt1 and "vectors" in st.session_state:
        document_chain = create_stuff_documents_chain(llm, prompt)
        retriever = st.session_state.vectors.as_retriever()
        retrieval_chain = create_retrieval_chain(retriever, document_chain)
        response = retrieval_chain.invoke({'input': prompt1})
        st.write(response['answer'])
